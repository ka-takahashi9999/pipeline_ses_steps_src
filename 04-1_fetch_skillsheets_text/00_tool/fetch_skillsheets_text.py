#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
04-1_fetch_skillsheets_text: 要員メールからスキルシートテキストを取得

入力:
  - 01-1_fetch_gmail/01_result/fetch_gmail_mail_master.jsonl  (添付ファイル参照)
  - 01-4_cleanup_email_text/01_result/cleanup_email_text_emails_raw.jsonl  (本文URL抽出)
  - 02-2_classify_output_file_project_resource/01_result/resources.jsonl  (処理対象)

出力:
  - 04-1_fetch_skillsheets_text/01_result/fetch_skillsheets_text.jsonl
  - 04-1_fetch_skillsheets_text/01_result/99_no_fetch_skillsheets_text.jsonl  (success=falseのみ)

取得優先順位:
  1. 添付ファイル (mail_master.jsonlのattachmentsを参照)
  2. URLから取得 (Google Drive > OneDrive > その他HTTP)
  3. すべて失敗時は success=false

LLM使用禁止
"""

import os
import sys
import re
import json
import base64
import shutil
import subprocess
import tempfile
import time
import html
from dataclasses import dataclass
import requests
import warnings
import logging
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple
from urllib.parse import parse_qs, urljoin, urlparse

# common モジュールのパスを追加
STEP_DIR = Path(__file__).resolve().parent.parent
PROJECT_ROOT = STEP_DIR.parent
sys.path.insert(0, str(PROJECT_ROOT))

from common.logger import get_logger
from common.json_utils import read_jsonl_as_list, append_jsonl
from common.file_utils import write_error_log, write_execution_time

logger = get_logger("04-1_fetch_skillsheets_text")

# ライブラリ警告を抑制
warnings.filterwarnings("ignore", category=UserWarning, module="openpyxl")
logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("pdfplumber").setLevel(logging.ERROR)

# ---------------------------------------------------------------------------
# 定数・パス
# ---------------------------------------------------------------------------
RESULT_DIR = STEP_DIR / "01_result"
EXEC_TIME_DIR = STEP_DIR / "99_execution_time"
NO_FETCH_JSONL = RESULT_DIR / "99_no_fetch_skillsheets_text.jsonl"
OUTPUT_JSONL = RESULT_DIR / "fetch_skillsheets_text.jsonl"

INPUT_MAIL_MASTER = PROJECT_ROOT / "01-1_fetch_gmail/01_result/fetch_gmail_mail_master.jsonl"
INPUT_CLEANUP_EMAILS = (
    PROJECT_ROOT
    / "01-4_cleanup_email_text/01_result/cleanup_email_text_emails_raw.jsonl"
)
INPUT_RESOURCES = (
    PROJECT_ROOT
    / "02-2_classify_output_file_project_resource/01_result/resources.jsonl"
)

# OneDrive ドメイン判定
ONEDRIVE_DOMAINS = {"onedrive.live.com", "1drv.ms", "sharepoint.com"}
CTSU_PUBLIC_DOMAINS = {"ctsu.jp", "cho-tatsu.com"}

# スキルシートらしいファイル拡張子
SKILLSHEET_EXTENSIONS = {"pdf", "xlsx", "xls", "docx", "doc"}

SHORTENER_DOMAINS = {
    "bit.ly",
    "tinyurl.com",
    "goo.gl",
    "ow.ly",
    "x.gd",
    "t.co",
    "buff.ly",
    "is.gd",
    "cutt.ly",
    "lnkd.in",
}

CONDITIONAL_HOST_DOMAINS = {
    "dropbox.com",
    "box.com",
    "d.bmb.jp",
    "a.bme.jp",
}

DIRECT_ADOPT_NG_RULES = (
    ("hm-f.jp", "/cc.php"),
)

TITLE_ONLY_TEXTS = {
    "スキルシート",
    "技術者経歴書",
    "経歴書",
    "職務経歴書",
}

# HTTP ダウンロード共通ヘッダー
HTTP_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    )
}

# タイムアウト設定
TIMEOUT_GDRIVE = 30   # Google Drive はやや大きめ
TIMEOUT_OTHER = 10    # その他 URL は短め
MIN_TEXT_LENGTH = 30


@dataclass
class DownloadResult:
    content: bytes
    final_url: str
    content_type: str
    filename: str = ""


def _base64url_encode(value: str) -> str:
    """URL-safe base64 をパディングなしで返す。"""
    return base64.urlsafe_b64encode(value.encode("utf-8")).decode("ascii").rstrip("=")

# ---------------------------------------------------------------------------
# テキスト抽出
# ---------------------------------------------------------------------------


def extract_text_from_pdf(data: bytes) -> str:
    """PDFバイナリからテキストを抽出する。"""
    import pdfplumber

    texts: List[str] = []
    with pdfplumber.open(BytesIO(data)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                texts.append(t)
    return "\n".join(texts)


def extract_text_from_excel(data: bytes, filename: str = "") -> str:
    """Excel (.xlsx / .xls) バイナリからテキストを抽出する。"""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # .xls (古い形式)
    if ext == "xls" and data[:4] != b"PK\x03\x04":
        import xlrd

        wb = xlrd.open_workbook(file_contents=data)
        rows: List[str] = []
        for sname in wb.sheet_names():
            sheet = wb.sheet_by_name(sname)
            rows.append(f"=== シート: {sname} ===")
            for ri in range(sheet.nrows):
                cells = [
                    str(sheet.cell(ri, ci).value).strip()
                    for ci in range(sheet.ncols)
                    if sheet.cell(ri, ci).value
                ]
                if cells:
                    rows.append(" | ".join(cells))
        return "\n".join(rows)

    # .xlsx (ZIP形式)
    import openpyxl

    wb = openpyxl.load_workbook(BytesIO(data), data_only=True)
    rows = []
    for sname in wb.sheetnames:
        sheet = wb[sname]
        rows.append(f"=== シート: {sname} ===")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                rows.append(" | ".join(cells))
    return "\n".join(rows)


def extract_text_from_word(data: bytes) -> str:
    """Word (.docx) バイナリからテキストを抽出する。"""
    import docx

    doc = docx.Document(BytesIO(data))
    parts: List[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _list_ole_streams(data: bytes) -> List[str]:
    """OLE2 コンテナのストリーム名一覧を返す。"""
    import olefile

    with olefile.OleFileIO(BytesIO(data)) as ole:
        return [entry[-1] for entry in ole.listdir()]


def detect_ole_office_type(data: bytes, filename: str = "") -> Tuple[Optional[str], str]:
    """
    OLE2 Office バイナリを Word / Excel に判定する。
    戻り値: ("word" | "excel" | None, detail)
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "doc":
        return "word", "拡張子 .doc"
    if ext == "xls":
        return "excel", "拡張子 .xls"

    try:
        streams = _list_ole_streams(data)
    except Exception as e:
        return None, f"OLEストリーム確認失敗: {e}"

    lowered = {name.lower() for name in streams}
    if "worddocument" in lowered:
        return "word", "OLEストリーム WordDocument"
    if "workbook" in lowered or "book" in lowered:
        return "excel", "OLEストリーム Workbook/Book"
    return None, f"OLEストリーム種別不明: {streams}"


def _run_legacy_word_command(args: List[str], data: bytes, suffix: str = ".doc") -> str:
    """旧Word抽出コマンドを一時ファイル経由で実行する。"""
    fd, path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    try:
        with open(path, "wb") as wf:
            wf.write(data)
        completed = subprocess.run(
            args + [path],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            check=False,
        )
        if completed.returncode != 0:
            stderr = (completed.stderr or "").strip()
            raise ValueError(stderr or f"returncode={completed.returncode}")
        return completed.stdout
    finally:
        try:
            os.remove(path)
        except OSError:
            pass


def extract_text_from_legacy_word(data: bytes, filename: str = "") -> str:
    """旧Word (.doc) バイナリからテキストを抽出する。"""
    command_candidates = [
        ("antiword", ["antiword"]),
        ("catdoc", ["catdoc"]),
        ("wvText", ["wvText"]),
    ]
    for label, args in command_candidates:
        if shutil.which(args[0]):
            try:
                text = _run_legacy_word_command(args, data, suffix=".doc")
                if text.strip():
                    return text
                raise ValueError("抽出テキストが空")
            except Exception as e:
                raise ValueError(f"旧Word(.doc)抽出失敗[{label}]: {e}") from e

    if shutil.which("strings"):
        try:
            text = _run_legacy_word_command(["strings", "-e", "l", "-n", "4"], data, suffix=".doc")
            cleaned_lines = []
            skip_terms = {
                "Root Entry",
                "WordDocument",
                "1Table",
                "SummaryInformation",
                "DocumentSummaryInformation",
                "CompObj",
            }
            for line in text.splitlines():
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped in skip_terms:
                    continue
                if re.fullmatch(r"0{4,}", stripped):
                    continue
                cleaned_lines.append(stripped)
            cleaned = "\n".join(cleaned_lines)
            if cleaned.strip():
                return cleaned
            raise ValueError("抽出テキストが空")
        except Exception as e:
            raise ValueError(f"旧Word(.doc)抽出失敗[strings]: {e}") from e

    raise ValueError("旧Word(.doc)抽出未対応: 利用可能コマンドなし")


def extract_text_from_bytes(data: bytes, filename: str = "") -> str:
    """ファイル形式を自動判定してテキストを抽出する。"""
    header = data[:4]
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if header == b"%PDF":
        return extract_text_from_pdf(data)

    if header == b"PK\x03\x04":
        # ZIP: Excel か Word か拡張子で判定
        if ext in ("xlsx", "xls"):
            return extract_text_from_excel(data, filename)
        if ext == "docx":
            return extract_text_from_word(data)
        # 拡張子不明 → Excel を先に試す
        try:
            return extract_text_from_excel(data, filename)
        except Exception:
            return extract_text_from_word(data)

    # OLE2 Office (旧Word / 旧Excel)
    if header == b"\xd0\xcf\x11\xe0":
        ole_type, detail = detect_ole_office_type(data, filename)
        if ole_type == "word":
            return extract_text_from_legacy_word(data, filename)
        if ole_type == "excel":
            return extract_text_from_excel(data, filename)
        raise ValueError(f"OLE2形式判定失敗: {detail}")

    if ext == "doc":
        return extract_text_from_legacy_word(data, filename)

    if ext == "xls":
        return extract_text_from_excel(data, filename)

    # テキストとして読む
    try:
        decoded = data.decode("utf-8")
        if decoded.strip():
            return decoded
    except UnicodeDecodeError:
        pass
    try:
        decoded = data.decode("cp932")
        if decoded.strip():
            return decoded
    except UnicodeDecodeError:
        pass

    raise ValueError(f"サポートされていないファイル形式: filename={filename!r}, header={header!r}")


# ---------------------------------------------------------------------------
# 添付ファイル処理
# ---------------------------------------------------------------------------


def extract_from_attachment(attachment: Dict[str, Any]) -> str:
    """
    Gmail添付ファイル辞書からテキストを抽出する。
    data フィールドは base64url エンコード済み。
    """
    data_b64 = attachment.get("data", "")
    if not data_b64:
        raise ValueError("attachment.data が空です")

    # base64url デコード（パディング調整）
    padded = data_b64 + "=" * (-len(data_b64) % 4)
    raw_bytes = base64.urlsafe_b64decode(padded)

    filename = attachment.get("filename", "")
    return extract_text_from_bytes(raw_bytes, filename)


# ---------------------------------------------------------------------------
# URL 抽出・分類
# ---------------------------------------------------------------------------


def extract_urls_from_text(text: str) -> List[str]:
    """テキスト中の https:// を含む URL をすべて抽出する。"""
    pattern = r'https?://[^\s\)\]\}\"\'<>]+'
    raw_urls = re.findall(pattern, text)
    # 末尾の不要な記号をトリム
    trailing = set(")]}>\"'.,;:!?")
    cleaned = []
    for u in raw_urls:
        while u and u[-1] in trailing:
            u = u[:-1]
        if u:
            cleaned.append(u)
    return cleaned


def classify_url(url: str) -> str:
    """
    URL を source 種別に分類する。
    戻り値: "google_drive" | "one_drive" | "ctsu_public_talent" | "other"
    """
    lower = url.lower()
    if "drive.google.com" in lower or "docs.google.com" in lower:
        return "google_drive"
    parsed = urlparse(url)
    host = parsed.netloc.lower()
    if is_ctsu_public_talent_url(url):
        return "ctsu_public_talent"
    for d in ONEDRIVE_DOMAINS:
        if host_matches(host, d):
            return "one_drive"
    return "other"


def host_matches(host: str, domain: str) -> bool:
    """ホストがドメインと一致またはサブドメインかを判定する。"""
    return host == domain or host.endswith(f".{domain}")


def get_url_extension(url: str) -> str:
    """URLパスからファイル拡張子を抽出する。"""
    path = urlparse(url).path.lower()
    if "." not in path.rsplit("/", 1)[-1]:
        return ""
    return path.rsplit(".", 1)[-1]


def is_google_docs_url(url: str) -> bool:
    """Google Drive / Docs のURLかを判定する。"""
    lower = url.lower()
    return "drive.google.com/file/d/" in lower or "docs.google.com/" in lower


def is_cloudinary_raw_url(url: str) -> bool:
    """Cloudinary raw/upload URL かを判定する。"""
    lower = url.lower()
    parsed = urlparse(url)
    return host_matches(parsed.netloc.lower(), "res.cloudinary.com") and "/raw/upload/" in lower


def is_direct_adopt_ng_url(url: str) -> bool:
    """直接採用してはいけないURLかを判定する。"""
    parsed = urlparse(url)
    host = parsed.netloc.lower()
    path = parsed.path.lower()
    for domain, path_prefix in DIRECT_ADOPT_NG_RULES:
        if host_matches(host, domain) and path.startswith(path_prefix):
            return True
    return False


def is_ctsu_public_talent_url(url: str) -> bool:
    """CTSU公開人材ページURLかを判定する。"""
    parsed = urlparse(url)
    host = parsed.netloc.lower()
    path = parsed.path.lower()
    return any(host_matches(host, domain) for domain in CTSU_PUBLIC_DOMAINS) and path.startswith("/public-talents/")


def classify_url_candidate(url: str) -> str:
    """
    URL候補を分類する。
    戻り値:
      - google_drive
      - one_drive
      - ctsu_public_talent
      - cloudinary_raw
      - extension_file
      - conditional
      - direct_adopt_ng
    """
    parsed = urlparse(url)
    host = parsed.netloc.lower()

    if is_google_docs_url(url):
        return "google_drive"

    for d in ONEDRIVE_DOMAINS:
        if host_matches(host, d):
            return "one_drive"

    if is_ctsu_public_talent_url(url):
        return "ctsu_public_talent"

    if is_cloudinary_raw_url(url):
        return "cloudinary_raw"

    if get_url_extension(url) in SKILLSHEET_EXTENSIONS:
        return "extension_file"

    if is_direct_adopt_ng_url(url):
        return "direct_adopt_ng"

    for d in SHORTENER_DOMAINS | CONDITIONAL_HOST_DOMAINS:
        if host_matches(host, d):
            return "conditional"

    if host:
        return "conditional"

    return "direct_adopt_ng"


def sort_urls_by_priority(urls: List[str]) -> List[Tuple[str, str]]:
    """
    URL リストを (url, category) のタプルリストにして優先順位順に並べ替える。
    優先:
      google_drive > one_drive > ctsu_public_talent > cloudinary_raw > extension_file > conditional > direct_adopt_ng
    """
    priority = {
        "google_drive": 0,
        "one_drive": 1,
        "ctsu_public_talent": 2,
        "cloudinary_raw": 3,
        "extension_file": 4,
        "conditional": 5,
        "direct_adopt_ng": 6,
    }
    tagged = [(u, classify_url_candidate(u)) for u in urls]
    tagged.sort(key=lambda x: priority[x[1]])
    # 重複 URL を除去（同じ URL が複数あっても1回だけ試す）
    seen: Set[str] = set()
    deduped = []
    for u, s in tagged:
        if u not in seen:
            seen.add(u)
            deduped.append((u, s))
    return deduped


def is_probable_file_url(url: str) -> bool:
    """URL文字列から実ファイルURLらしさを判定する。"""
    return (
        is_google_docs_url(url)
        or is_cloudinary_raw_url(url)
        or get_url_extension(url) in SKILLSHEET_EXTENSIONS
    )


def is_probable_file_response(final_url: str, content_type: str) -> bool:
    """最終URLまたはContent-Typeから実ファイル応答らしさを判定する。"""
    normalized_type = (content_type or "").lower()
    if is_probable_file_url(final_url):
        return True
    file_markers = (
        "application/pdf",
        "application/msword",
        "application/vnd.ms-excel",
        "application/vnd.openxmlformats-officedocument",
        "application/octet-stream",
        "binary/octet-stream",
    )
    return any(marker in normalized_type for marker in file_markers)


def is_html_response(content: bytes, content_type: str) -> bool:
    """HTMLページ応答かを緩めに判定する。"""
    normalized_type = (content_type or "").lower()
    if "text/html" in normalized_type or "application/xhtml+xml" in normalized_type:
        return True
    head = content[:1000].lower()
    html_markers = (b"<!doctype html", b"<html", b"<head", b"<body")
    return any(marker in head for marker in html_markers)


def extract_html_redirect_url(content: bytes, base_url: str) -> Optional[str]:
    """HTML本文中の JS / meta refresh リダイレクト URL を抽出する。"""
    try:
        html = content[:20000].decode("utf-8", errors="ignore")
    except Exception:
        return None

    patterns = (
        r'http-equiv=["\']refresh["\'][^>]*content=["\'][^"\']*url=([^"\'>]+)',
        r'location(?:\.href|\.replace)?\s*=\s*["\']([^"\']+)["\']',
        r'window\.open\(\s*["\']([^"\']+)["\']',
    )
    for pattern in patterns:
        match = re.search(pattern, html, flags=re.IGNORECASE)
        if match:
            return urljoin(base_url, match.group(1).strip())
    return None


def validate_skillsheet_text(text: str) -> Optional[str]:
    """success=true にする前の最低限の品質チェック。問題がなければ None を返す。"""
    stripped = text.strip()
    if not stripped:
        return "抽出テキストが空"

    lowered = stripped.lower()
    if any(marker in lowered for marker in ("<!doctype html", "<html", "<head", "<body")):
        return "HTML本文のため不採用"

    compact = re.sub(r"\s+", "", stripped)
    if len(compact) < MIN_TEXT_LENGTH:
        return f"テキストが短すぎるため不採用({len(compact)}文字)"

    lines = [line.strip() for line in stripped.splitlines() if line.strip()]
    if lines:
        page_number_lines = [
            line for line in lines
            if re.fullmatch(r"\d+\s*/\s*\d*", line)
        ]
        page_number_ratio = len(page_number_lines) / len(lines)
        if page_number_lines and (
            len(page_number_lines) == len(lines)
            or (len(lines) >= 10 and page_number_ratio >= 0.9)
        ):
            return "ページ番号だけの本文のため不採用"

    if lines and len(lines) <= 3:
        normalized = "".join(lines)
        normalized = re.sub(r"[ 　\t\r\n:：\-_/()（）【】\[\]]+", "", normalized)
        if normalized in TITLE_ONLY_TEXTS:
            return "表題だけの本文のため不採用"

    return None


def classify_attachment_failure(filename: str, error: Exception) -> str:
    """添付ファイル失敗理由を対象区分つきで整形する。"""
    message = str(error).strip() or "unknown attachment error"
    if (
        "could not read strings" in message.lower()
        or "サポートされていないファイル形式" in message
        or "旧Word(.doc)抽出未対応" in message
        or "抽出テキストが空" in message
    ):
        scope = "対象外"
        kind = "extract_unsupported_file"
    else:
        scope = "要確認"
        kind = "attachment_fetch_failed"
    return f"{scope}[{kind}] attachment:{filename}: {message}"


def classify_url_failure(category: str, url: str, error: Exception) -> str:
    """URL失敗理由を対象区分つきで整形する。"""
    message = str(error).strip() or "unknown url error"
    host = urlparse(url).netloc.lower()
    lower = message.lower()

    if host_matches(host, "d.bmb.jp") or host_matches(host, "a.bme.jp"):
        scope = "改善対象"
        kind = "redirect_url_resolution"
    elif category == "one_drive" or host_matches(host, "1drv.ms"):
        scope = "改善対象"
        kind = "onedrive_short_url_resolution"
    elif category == "ctsu_public_talent":
        scope = "改善対象"
        kind = "ctsu_fileid_resolution"
    elif any(code in lower for code in ("httpエラー 401", "httpエラー 403", "httpエラー 404", "httpエラー 410")):
        scope = "対象外"
        kind = "expired_or_permission_denied"
    elif "htmlページ応答のため不採用" in lower:
        scope = "対象外"
        kind = "no_actual_file_url"
    elif (
        "抽出テキストが空" in message
        or "サポートされていないファイル形式" in message
        or "テキストが短すぎるため不採用" in message
        or "表題だけの本文のため不採用" in message
    ):
        scope = "対象外"
        kind = "extract_unsupported_file"
    else:
        scope = "要確認"
        kind = "url_fetch_failed"
    return f"{scope}[{kind}] url[{category}] {url}: {message}"


def resolve_redirect_url(
    session: requests.Session,
    url: str,
    timeout: int = TIMEOUT_OTHER,
    max_hops: int = 8,
) -> str:
    """Location / meta refresh / JS redirect を手動追跡して最終URL候補を返す。"""
    current_url = url
    for _ in range(max_hops):
        response = session.get(
            current_url,
            headers=HTTP_HEADERS,
            timeout=timeout,
            allow_redirects=False,
            stream=True,
        )

        if 300 <= response.status_code < 400:
            location = (response.headers.get("Location") or "").strip()
            if not location:
                raise ValueError(f"HTTPリダイレクト先が空です: {current_url}")
            current_url = urljoin(current_url, location)
            continue

        if response.status_code == 200:
            if is_html_response(response.content, response.headers.get("Content-Type", "")):
                redirect_url = extract_html_redirect_url(
                    response.content,
                    response.url or current_url,
                )
                if redirect_url and redirect_url != current_url:
                    current_url = redirect_url
                    continue
            return response.url or current_url

        raise ValueError(f"HTTPエラー {response.status_code}: {current_url}")

    raise ValueError(f"リダイレクト追跡上限超過: {url}")


# ---------------------------------------------------------------------------
# Google Drive ダウンロード
# ---------------------------------------------------------------------------


def extract_gdrive_file_id(url: str) -> Optional[str]:
    """Google Drive URL からファイル ID を抽出する。"""
    patterns = [
        r"open\?id=([^&\s]+)",
        r"uc\?.*id=([^&\s]+)",
        r"drive\.google\.com/file/d/([^/\?\s]+)",
        r"docs\.google\.com/spreadsheets/d/([^/\?\s]+)",
        r"docs\.google\.com/document/d/([^/\?\s]+)",
        r"/d/([a-zA-Z0-9_-]{20,})",
    ]
    for p in patterns:
        m = re.search(p, url)
        if m:
            return m.group(1)
    return None


def _detect_google_doc_type(url: str) -> str:
    """
    Google Docs URL の種別を判定する。
    戻り値: "spreadsheet" | "document" | "file"
    """
    lower = url.lower()
    if "docs.google.com/spreadsheets/" in lower:
        return "spreadsheet"
    if "docs.google.com/document/" in lower:
        return "document"
    return "file"


def download_google_drive(url: str, timeout: int = TIMEOUT_GDRIVE) -> DownloadResult:
    """
    Google Drive URL からファイルをダウンロードして bytes を返す。
    Spreadsheet は xlsx、Document は docx で export する。
    失敗時は例外を送出する。
    """
    # open?id= 形式を正規化
    parsed = urlparse(url)
    if parsed.path in ("/open", "/uc"):
        params = parse_qs(parsed.query)
        file_id = params.get("id", [None])[0]
    else:
        file_id = extract_gdrive_file_id(url)

    if not file_id:
        raise ValueError(f"Google Drive ファイルIDを取得できません: {url}")

    doc_type = _detect_google_doc_type(url)
    session = requests.Session()

    # Spreadsheet / Document は export API を使う
    if doc_type == "spreadsheet":
        download_url = f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=xlsx"
    elif doc_type == "document":
        download_url = f"https://docs.google.com/document/d/{file_id}/export?format=docx"
    else:
        download_url = f"https://drive.google.com/uc?export=download&id={file_id}"

    resp = session.get(download_url, stream=True, timeout=timeout, headers=HTTP_HEADERS, allow_redirects=True)

    # uc?export=download の場合のみウイルス警告クッキー処理
    if doc_type == "file":
        for k, v in resp.cookies.items():
            if k.startswith("download_warning"):
                resp = session.get(
                    download_url,
                    params={"id": file_id, "confirm": v},
                    stream=True,
                    timeout=timeout,
                    headers=HTTP_HEADERS,
                )
                break

    if resp.status_code != 200:
        raise ValueError(f"HTTPエラー {resp.status_code}: {download_url}")

    content = resp.content
    if not content:
        raise ValueError("ダウンロードされたコンテンツが空です")

    # HTML エラーページ判定
    if len(content) < 10000 and b"<html" in content[:200]:
        raise ValueError("Google Drive がエラーHTMLを返しました (認証が必要か公開設定されていない可能性)")

    return DownloadResult(
        content=content,
        final_url=resp.url or download_url,
        content_type=resp.headers.get("Content-Type", ""),
        filename="",
    )


# ---------------------------------------------------------------------------
# OneDrive ダウンロード
# ---------------------------------------------------------------------------


def download_onedrive(url: str, timeout: int = TIMEOUT_OTHER) -> DownloadResult:
    """
    OneDrive URL からファイルをダウンロードして bytes を返す。
    1drv.ms 短縮 URL のリダイレクトも追跡する。
    """
    session = requests.Session()

    actual_url = url
    if "1drv.ms" in url:
        try:
            actual_url = resolve_redirect_url(session, url, timeout=timeout)
        except Exception:
            actual_url = url

    share_candidates: List[str] = []
    for candidate in (url, actual_url):
        if not candidate:
            continue
        token = _base64url_encode(candidate)
        share_candidates.append(
            f"https://api.onedrive.com/v1.0/shares/u!{token}/root/content"
        )

    for candidate_url in share_candidates:
        try:
            resp = session.get(
                candidate_url,
                stream=True,
                timeout=timeout,
                headers=HTTP_HEADERS,
                allow_redirects=True,
            )
        except Exception:
            continue
        if resp.status_code != 200 or not resp.content:
            continue
        if is_html_response(resp.content, resp.headers.get("Content-Type", "")):
            continue
        return DownloadResult(
            content=resp.content,
            final_url=resp.url or candidate_url,
            content_type=resp.headers.get("Content-Type", ""),
            filename="",
        )

    # OneDrive の共有リンクはダウンロード URL に変換
    # https://onedrive.live.com/... → download パラメータ追加
    if "onedrive.live.com" in actual_url and "download" not in actual_url.lower():
        if "?" in actual_url:
            actual_url += "&download=1"
        else:
            actual_url += "?download=1"

    resp = session.get(actual_url, stream=True, timeout=timeout, headers=HTTP_HEADERS, allow_redirects=True)
    if resp.status_code != 200:
        raise ValueError(f"HTTPエラー {resp.status_code}: {actual_url}")

    content = resp.content
    if not content:
        raise ValueError("ダウンロードされたコンテンツが空です")

    return DownloadResult(
        content=content,
        final_url=resp.url or actual_url,
        content_type=resp.headers.get("Content-Type", ""),
        filename="",
    )


# ---------------------------------------------------------------------------
# その他 URL ダウンロード
# ---------------------------------------------------------------------------


def download_other_url(url: str, timeout: int = TIMEOUT_OTHER) -> DownloadResult:
    """任意の URL からファイルをダウンロードして bytes を返す。
    リダイレクト先が Google Docs/Drive の場合は専用処理に委譲する。
    """
    session = requests.Session()

    try:
        final_url = resolve_redirect_url(session, url, timeout=timeout)
    except Exception:
        final_url = url

    # リダイレクト先が Google Docs/Drive なら専用処理に委譲
    final_lower = final_url.lower()
    if "drive.google.com" in final_lower or "docs.google.com" in final_lower:
        return download_google_drive(final_url, timeout=TIMEOUT_GDRIVE)
    final_host = urlparse(final_lower).netloc.lower()
    if any(host_matches(final_host, d) for d in ONEDRIVE_DOMAINS):
        return download_onedrive(final_url, timeout=TIMEOUT_OTHER)

    # 通常のダウンロード
    resp = session.get(
        final_url,
        stream=True,
        timeout=timeout,
        headers=HTTP_HEADERS,
        allow_redirects=True,
    )
    if resp.status_code != 200:
        raise ValueError(f"HTTPエラー {resp.status_code}: {final_url}")

    content = resp.content
    if not content:
        raise ValueError("ダウンロードされたコンテンツが空です")

    return DownloadResult(
        content=content,
        final_url=resp.url or final_url,
        content_type=resp.headers.get("Content-Type", ""),
        filename="",
    )


def _extract_ctsu_file_ids(html: str) -> List[str]:
    """CTSU公開ページHTMLから fileId 候補を抽出する。"""
    patterns = (
        r'fileId\s*[:=]\s*["\']?([A-Za-z0-9_-]{8,})',
        r'"fileId"\s*:\s*"([A-Za-z0-9_-]{8,})"',
        r'"fileId"\s*,\s*"([A-Za-z0-9_-]{8,})"',
        r'fileId["\']?\s*,\s*["\']([A-Za-z0-9_-]{8,})["\']',
        r'\\"fileId\\"\s*,\s*\\"([A-Za-z0-9_-]{8,})\\"',
        r'/api/files/([A-Za-z0-9_-]{8,})',
    )
    found: List[str] = []
    seen: Set[str] = set()
    for pattern in patterns:
        for file_id in re.findall(pattern, html, flags=re.IGNORECASE):
            if file_id not in seen:
                seen.add(file_id)
                found.append(file_id)
    return found


def _extract_text_from_html_document(raw_html: str) -> str:
    """HTML文書から可読テキストだけを抜き出す。"""
    cleaned = re.sub(r"(?is)<script[^>]*>.*?</script>", " ", raw_html)
    cleaned = re.sub(r"(?is)<style[^>]*>.*?</style>", " ", cleaned)
    cleaned = re.sub(r"(?is)<noscript[^>]*>.*?</noscript>", " ", cleaned)
    cleaned = re.sub(r"(?i)<br\\s*/?>", "\n", cleaned)
    cleaned = re.sub(r"(?i)</(?:p|div|section|article|li|tr|h[1-6])>", "\n", cleaned)
    cleaned = re.sub(r"(?s)<[^>]+>", " ", cleaned)
    cleaned = html.unescape(cleaned)
    cleaned = cleaned.replace("\r", "\n")
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    cleaned = re.sub(r"[ \t\f\v]+", " ", cleaned)
    return "\n".join(line.strip() for line in cleaned.splitlines() if line.strip())


def _normalize_ctsu_file_payload(payload: Any) -> Dict[str, Any]:
    """CTSU /api/files 応答から signedUrl/fileName を持つ辞書を取り出す。"""
    if isinstance(payload, dict):
        if payload.get("signedUrl"):
            return payload
        for key in ("data", "file", "result"):
            child = payload.get(key)
            normalized = _normalize_ctsu_file_payload(child)
            if normalized:
                return normalized
    elif isinstance(payload, list):
        for item in payload:
            normalized = _normalize_ctsu_file_payload(item)
            if normalized:
                return normalized
    return {}


def _score_ctsu_candidate(file_name: str, signed_url: str) -> Tuple[int, int]:
    """CTSU候補の優先度スコアを返す。大きいほど優先。"""
    base_name = (file_name or urlparse(signed_url).path.rsplit("/", 1)[-1]).lower()
    ext = base_name.rsplit(".", 1)[-1] if "." in base_name else ""
    ext_priority = {
        "pdf": 5,
        "docx": 4,
        "doc": 3,
        "xlsx": 2,
        "xls": 1,
    }.get(ext, 0)
    keyword_score = 0
    for kw in ("スキルシート", "経歴書", "職務経歴書"):
        if kw in (file_name or ""):
            keyword_score += 1
    return keyword_score, ext_priority


def _select_ctsu_candidate(candidates: List[Dict[str, str]]) -> Dict[str, str]:
    """CTSU候補一覧から優先順位に従って1件選ぶ。"""
    if not candidates:
        return {}
    ranked = sorted(
        enumerate(candidates),
        key=lambda item: (_score_ctsu_candidate(item[1].get("file_name", ""), item[1].get("signed_url", "")), -item[0]),
        reverse=True,
    )
    best_score = _score_ctsu_candidate(ranked[0][1].get("file_name", ""), ranked[0][1].get("signed_url", ""))
    if best_score == (0, 0):
        return candidates[0]
    return ranked[0][1]


def download_ctsu_public_talent(url: str, timeout: int = TIMEOUT_OTHER) -> DownloadResult:
    """
    CTSU公開人材ページから fileId -> /api/files/{fileId} -> signedUrl の順で実ファイルを取得する。
    """
    session = requests.Session()

    try:
        page_resp = session.get(
            url,
            timeout=timeout,
            headers=HTTP_HEADERS,
            allow_redirects=True,
        )
    except Exception as e:
        raise ValueError(f"CTSU公開ページ取得失敗: {e}") from e

    if page_resp.status_code != 200:
        raise ValueError(f"CTSU公開ページ取得失敗: HTTP {page_resp.status_code}")

    html = page_resp.text or ""
    file_ids = _extract_ctsu_file_ids(html)
    if not file_ids:
        page_text = _extract_text_from_html_document(html)
        if validate_skillsheet_text(page_text) is None:
            return DownloadResult(
                content=page_text.encode("utf-8"),
                final_url=page_resp.url or url,
                content_type="text/plain; charset=utf-8",
                filename="ctsu_public_talent.txt",
            )
        raise ValueError("CTSU fileId 抽出失敗")

    candidates: List[Dict[str, str]] = []
    api_errors: List[str] = []
    for file_id in file_ids:
        api_url = f"https://ctsu.jp/api/files/{file_id}"
        try:
            api_resp = session.get(
                api_url,
                timeout=timeout,
                headers=HTTP_HEADERS,
                allow_redirects=True,
            )
        except Exception as e:
            api_errors.append(f"{file_id}: {e}")
            continue

        if api_resp.status_code != 200:
            api_errors.append(f"{file_id}: HTTP {api_resp.status_code}")
            continue

        try:
            payload = api_resp.json()
        except json.JSONDecodeError as e:
            api_errors.append(f"{file_id}: JSON decode error {e}")
            continue

        file_info = _normalize_ctsu_file_payload(payload)
        signed_url = file_info.get("signedUrl")
        file_name = file_info.get("fileName", "")
        if not signed_url:
            api_errors.append(f"{file_id}: signedUrl missing")
            continue

        candidates.append(
            {
                "file_id": file_id,
                "signed_url": signed_url,
                "file_name": file_name,
            }
        )

    if not candidates:
        detail = ", ".join(api_errors[:3]) if api_errors else "candidate missing"
        raise ValueError(f"CTSU api/files 応答不正: {detail}")

    selected = _select_ctsu_candidate(candidates)
    signed_url = selected["signed_url"]

    try:
        file_resp = session.get(
            signed_url,
            timeout=timeout,
            headers=HTTP_HEADERS,
            allow_redirects=True,
        )
    except Exception as e:
        raise ValueError(f"CTSU signedUrl ダウンロード失敗: {e}") from e

    if file_resp.status_code != 200:
        raise ValueError(f"CTSU signedUrl ダウンロード失敗: HTTP {file_resp.status_code}")

    content = file_resp.content
    if not content:
        raise ValueError("CTSU signedUrl ダウンロード失敗: empty content")

    return DownloadResult(
        content=content,
        final_url=file_resp.url or signed_url,
        content_type=file_resp.headers.get("Content-Type", ""),
        filename=selected.get("file_name", ""),
    )


# ---------------------------------------------------------------------------
# メイン取得処理
# ---------------------------------------------------------------------------


def fetch_skillsheet(
    message_id: str,
    mail_record: Optional[Dict[str, Any]],
    cleanup_record: Optional[Dict[str, Any]],
) -> Dict[str, Any]:
    """
    スキルシートを取得してレコードを返す。

    Returns:
        {message_id, skillsheet, source, success, urls}
    """
    # --- ① 添付ファイル優先 ---
    attachments = (mail_record or {}).get("attachments") or []
    failure_reasons: List[str] = []
    if attachments:
        errors: List[str] = []
        classified_errors: List[str] = []
        for att in attachments:
            filename = att.get("filename", "")
            # スキルシートらしい添付ファイルのみ対象（PDF/Excel/Word）
            ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
            mime = att.get("mime_type", "").lower()
            if ext not in ("pdf", "xlsx", "xls", "docx", "doc") and not any(
                k in mime for k in ("pdf", "excel", "word", "spreadsheet", "document", "officedocument")
            ):
                # 添付がスキルシート対象外（画像等）の場合はスキップ
                continue
            try:
                text = extract_from_attachment(att)
                quality_error = validate_skillsheet_text(text)
                if quality_error is None:
                    logger.ok(f"添付ファイルからテキスト取得: {filename}", message_id)
                    return {
                        "message_id": message_id,
                        "skillsheet": text,
                        "source": "attachment",
                        "success": True,
                        "urls": False,
                    }
                errors.append(f"{filename}: {quality_error}")
                classified_errors.append(
                    classify_attachment_failure(filename, ValueError(quality_error))
                )
            except Exception as e:
                errors.append(f"{filename}: {e}")
                classified_errors.append(classify_attachment_failure(filename, e))
                logger.warn(f"添付ファイル抽出失敗 {filename}: {e}", message_id)

        # 添付はあったが全部失敗 or 対象外だった場合は URL に進む
        if errors:
            failure_reasons.extend(classified_errors)
            logger.warn(f"添付ファイル取得失敗（URL抽出に進む）: {errors}", message_id)

    # --- ② URL からの取得 ---
    body_text = (cleanup_record or {}).get("body_text", "") or ""
    urls = extract_urls_from_text(body_text)
    sorted_urls = sort_urls_by_priority(urls)

    tried_urls: List[str] = []
    for url, category in sorted_urls:
        tried_urls.append(url)
        source = classify_url(url)
        try:
            if source == "google_drive":
                download_result = download_google_drive(url)
            elif source == "one_drive":
                download_result = download_onedrive(url)
            elif source == "ctsu_public_talent":
                download_result = download_ctsu_public_talent(url)
            else:
                download_result = download_other_url(url)

            if category in ("conditional", "direct_adopt_ng"):
                if is_html_response(download_result.content, download_result.content_type):
                    raise ValueError("HTMLページ応答のため不採用")
                if not is_probable_file_response(download_result.final_url, download_result.content_type):
                    raise ValueError(
                        f"最終URLが実ファイルと判定できません: {download_result.final_url}"
                    )

            if is_html_response(download_result.content, download_result.content_type):
                raise ValueError("HTMLページ応答のため不採用")

            # URL のファイル名から拡張子推定
            if download_result.filename:
                url_filename = download_result.filename
            else:
                filename_url = download_result.final_url if is_probable_file_url(download_result.final_url) else url
                url_filename = urlparse(filename_url).path.rsplit("/", 1)[-1]
            text = extract_text_from_bytes(download_result.content, url_filename)

            quality_error = validate_skillsheet_text(text)
            if quality_error is None:
                logger.ok(
                    f"URLからテキスト取得 ({source}/{category}): {url} -> {download_result.final_url}",
                    message_id,
                )
                return {
                    "message_id": message_id,
                    "skillsheet": text,
                    "source": source,
                    "success": True,
                    "urls": url,
                }
            if source == "ctsu_public_talent" and quality_error == "抽出テキストが空":
                raise ValueError("CTSU 抽出テキストが空")
            raise ValueError(quality_error)
        except Exception as e:
            reason = classify_url_failure(category, url, e)
            failure_reasons.append(reason)
            logger.warn(f"URL取得失敗 ({source}/{category}) {url}: {e}", message_id)
            continue

    # --- ③ 取得失敗 ---
    urls_value = tried_urls[0] if tried_urls else False
    logger.warn(f"スキルシート取得失敗 tried_urls={tried_urls}", message_id)
    result = {
        "message_id": message_id,
        "skillsheet": None,
        "source": None,
        "success": False,
        "urls": urls_value,
    }
    if tried_urls:
        result["tried_urls"] = tried_urls
    if failure_reasons:
        result["failure_reason"] = " | ".join(failure_reasons[:10])
    else:
        if urls:
            result["failure_reason"] = "要確認[unknown_failure] URL取得失敗理由が記録されていません"
        else:
            result["failure_reason"] = "対象外[no_actual_file_url] 本文に実URLなし"
    return result


# ---------------------------------------------------------------------------
# エントリポイント
# ---------------------------------------------------------------------------


def load_mail_index(target_ids: Set[str]) -> Dict[str, Dict[str, Any]]:
    """
    mail_master.jsonl から対象 message_id のレコードのみ読み込む。
    添付データ(base64)を含むため全件読むとメモリが肥大する。
    """
    index: Dict[str, Dict[str, Any]] = {}
    logger.info(f"mail_master.jsonl 読み込み (対象IDのみ): {INPUT_MAIL_MASTER}")
    for rec in read_jsonl_as_list(str(INPUT_MAIL_MASTER)):
        mid = rec.get("message_id")
        if mid and mid in target_ids:
            index[mid] = rec
    logger.info(f"mail_master 読み込み完了: {len(index)} 件 (対象 {len(target_ids)} 件中)")
    return index


def main() -> None:
    start_time = time.time()

    # 出力ディレクトリ作成
    RESULT_DIR.mkdir(parents=True, exist_ok=True)
    EXEC_TIME_DIR.mkdir(parents=True, exist_ok=True)

    logger.info("=== 04-1_fetch_skillsheets_text 開始 ===")

    # 入力ファイル存在チェック
    for p in (INPUT_RESOURCES, INPUT_MAIL_MASTER, INPUT_CLEANUP_EMAILS):
        if not p.exists():
            err = FileNotFoundError(f"入力ファイルが見つかりません: {p}")
            logger.error(str(err))
            write_error_log(str(RESULT_DIR), err, context="入力ファイル確認")
            sys.exit(1)

    # --- 処理対象 ID 読み込み ---
    logger.info(f"resources.jsonl 読み込み: {INPUT_RESOURCES}")
    resources = read_jsonl_as_list(str(INPUT_RESOURCES))
    target_ids_list = [r["message_id"] for r in resources]
    target_ids_set = set(target_ids_list)
    total = len(target_ids_list)
    logger.info(f"処理対象: {total} 件")

    # --- mail_master: 対象IDのみロード（メモリ節約）---
    mail_index = load_mail_index(target_ids_set)

    # --- cleanup_emails: 対象IDのみロード ---
    logger.info(f"cleanup_emails.jsonl 読み込み: {INPUT_CLEANUP_EMAILS}")
    cleanup_index: Dict[str, Dict[str, Any]] = {}
    for rec in read_jsonl_as_list(str(INPUT_CLEANUP_EMAILS)):
        mid = rec.get("message_id")
        if mid and mid in target_ids_set:
            cleanup_index[mid] = rec
    logger.info(f"cleanup_emails 読み込み完了: {len(cleanup_index)} 件")

    # --- 出力ファイル初期化（都度 append するため事前クリア）---
    OUTPUT_JSONL.write_text("", encoding="utf-8")
    NO_FETCH_JSONL.write_text("", encoding="utf-8")

    # --- 処理（1件ずつ書き込み）---
    success_count = 0
    fail_count = 0

    for i, message_id in enumerate(target_ids_list, 1):
        mail_rec = mail_index.get(message_id)
        cleanup_rec = cleanup_index.get(message_id)

        record = fetch_skillsheet(message_id, mail_rec, cleanup_rec)

        # 即時書き込み（途中終了しても結果が残る）
        append_jsonl(str(OUTPUT_JSONL), record)
        if record["success"]:
            success_count += 1
        else:
            fail_count += 1
            append_jsonl(str(NO_FETCH_JSONL), record)

        # 10件ごとに進捗ログ
        if i % 10 == 0 or i == total:
            elapsed = time.time() - start_time
            rate = success_count / i * 100
            avg = elapsed / i
            logger.info(
                f"進捗: {i}/{total} | 成功:{success_count} 失敗:{fail_count} "
                f"取得率:{rate:.1f}% | 経過:{elapsed:.0f}秒 平均:{avg:.1f}秒/件"
            )

    # --- 最終サマリ ---
    elapsed = time.time() - start_time
    rate = success_count / total * 100 if total else 0.0

    logger.info("=== 完了 ===")
    logger.info(f"処理件数: {total} 件")
    logger.info(f"成功: {success_count} 件 / 失敗: {fail_count} 件 (取得率 {rate:.1f}%)")
    logger.info(f"出力: {OUTPUT_JSONL}")
    logger.info(f"未取得: {NO_FETCH_JSONL}")

    write_execution_time(
        str(EXEC_TIME_DIR),
        "fetch_skillsheets_text",
        elapsed,
        record_count=total,
    )

    if rate < 90.0:
        logger.warn(f"取得率が90%未満です: {rate:.1f}% (成功 {success_count}/{total})")

    sys.exit(0)


if __name__ == "__main__":
    main()
